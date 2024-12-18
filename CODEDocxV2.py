import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.scrolled import ScrolledText
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from tkinter import filedialog, messagebox
import tkinter as tk
from PIL import Image, ImageTk
import os
import re
import requests
from typing import List
import logging

class GoogleFontManager:
    """Manages downloading and loading of Google Fonts"""
    
    # Use environment variable for API key
    GOOGLE_FONTS_API = "https://www.googleapis.com/webfonts/v1/webfonts"
    
    DEFAULT_FONTS = [
        'Adobe Caslon Pro', 'Algerian', 'Andalus','Angsana New', 'AngsanaUPC', 'Arabic Typesetting', 'Arial', 'Arial Black', 'Arial Condensed', 'Arial Italic','Arial Narrow', 'Arial Narrow Bold', 'Arial Regular', 'Arial Rounded', 'Arial Rounded MT Bold','Arial Unicode', 'Arial Unicode MS', 'Baskerville','Baskerville Old Face', 'Batang', 'Bauhaus 93','Bell MT', 'Bernard MT Condensed', 'Blackadder ITC','Bodoni', 'Bodoni MT', 'Bodoni MT Condensed','Bodoni MT Poster', 'Book Antiqua', 'Bookman', 'Bookman Old Style','Brush Script', 'Brush Script MT', 'Calibri', 'Calibri Regular', 'Calisto','Calisto MT', 'Cambria', 'Cambria Math', 'Candara', 'Century', 'Century Gothic', 'Century Schoolbook', 'Comic Sans','Comic Sans Bold', 'Comic Sans MS', 'Comic Sans Regular','Consolas', 'Consolas Regular','Copperplate Gothic', 'Copperplate Gothic Bold', 'Copperplate Gothic Light','Corbel', 'Corbel Light','Corbel Regular', 'Courier', 'Courier Condensed','Courier New', 'Courier New Italic','Courier New Regular', 'Courier Regular', 'Cursive','Didot', 'Ebrima', 'Edwardian Script ITC','Felix Titling', 'Forte', 'Frank Ruhl Libre','Franklin Gothic Medium', 'Freestyle Script', 'Frutiger','Futura', 'Garamond', 'Garamond Bold', 'Garamond Condensed', 'Garamond Premier','Garamond Premier Pro', 'Garamond Pro', 'Georgia', 'Georgia Bold', 'Georgia Italic','Georgia Pro', 'Gill Sans','Gill Sans Regular', 'Goudy Old Style','Gulim', 'Gungsuh', 'GungsuhChe','Haettenschweiler', 'Harlow Solid Italic', 'Helvetica','Hoefler Text', 'Impact', 'Impact Italic', 'Ink Free','Jokerman', 'Kristen ITC', 'Lao UI','Lobster', 'Lucida Bright', 'Lucida Calligraphy', 'Lucida Console', 'Lucida Console Regular','Lucida Handwriting', 'Lucida Handwriting', 'Lucida Sans', 'Lucida Sans Italic', 'Lucida Sans Regular','Lucida Sans Unicode', 'MS Gothic', 'MS Mincho','MS PGothic', 'MS Reference Sans Serif', 'MS Reference Specialty','MS Sans Serif', 'MS Sans Serif Bold', 'MS Serif','MS UI Gothic', 'MV Boli', 'Magneto','Maiandra GD', 'Mangal', 'Marlett','Minion Pro', 'Mistral', 'Modern No. 20', 'Mongolian Baiti', 'Mongolian Regular', 'Monotype Corsiva', 'Myanmar Text', 'Myriad Pro','Nimbus Sans', 'OCR A Extended', 'Optima','Optima Italic', 'Palatino', 'Palatino Linotype', 'Papyrus', 'Papyrus Regular', 'Perpetua','Perpetua Titling', 'Perpetua Titling MT', 'Playbill','Poor Richard', 'Pristina', 'Quicksand','Rockwell', 'Rockwell Bold', 'Rockwell Condensed', 'Rockwell Extra Bold', 'Rockwell Light', 'Rockwell Regular', 'Rockwell Ultra Bold', 'Segoe Print', 'Segoe Script', 'Segoe UI', 'Segoe UI Black','Segoe UI Emoji', 'Segoe UI Historic', 'Segoe UI Light','Segoe UI Regular', 'Segoe UI Semibold', 'Segoe UI Symbol','Showcard', 'Showcard Gothic', 'SimHei','SimSun', 'Simplified Arabic', 'Simplified Arabic Fixed','Sitka', 'Snap ITC', 'Snap Regular', 'Swansea', 'Sylfaen','Tahoma', 'Tahoma Regular', 'Tempus Sans ITC', 'Times','Times New Roman', 'Times New Roman Bold', 'Times New Roman PS','Times New Roman Regular', 'Times Regular', 'Trebuchet','Trebuchet MS', 'Trebuchet Regular', 'Verdana', 'Verdana Bold', 'Verdana Italic','Verdana Regular', 'Viner Hand ITC', 'Wingdings', 'Wingdings 2', 'Wingdings 3','Yu Gothic', 'Yu Gothic UI', 'Zapf Dingbats','Zapfino', 'Zawgyi-One'


    ]
    
    @classmethod
    def get_font_list(cls, max_fonts: int = 50) -> List[str]:
        """
        Retrieve list of available Google Fonts with robust error handling
        
        Args:
            max_fonts (int): Maximum number of Google Fonts to retrieve
        
        Returns:
            List of font names, combining default fonts with Google Fonts
        """
        # Retrieve API key from environment variable
        api_key = os.getenv('GOOGLE_FONTS_API_KEY')
        if not api_key:
            logging.warning("Google Fonts API key not found. Falling back to default fonts.")
            return cls.DEFAULT_FONTS
        
        try:
            # Add API key as a parameter
            params = {'key': api_key}
            response = requests.get(cls.GOOGLE_FONTS_API, params=params)
            
            # Raise an exception for bad responses
            response.raise_for_status()
            
            # Extract font names
            fonts = response.json().get('items', [])
            google_fonts = [font['family'] for font in fonts[:max_fonts]]
            
            # Combine and deduplicate fonts
            combined_fonts = list(dict.fromkeys(cls.DEFAULT_FONTS + google_fonts))
            
            return combined_fonts
        
        except requests.RequestException as e:
            logging.error(f"Error fetching fonts: {e}")
            return cls.DEFAULT_FONTS
    
    @classmethod
    def set_api_key(cls, api_key: str):
        """
        Set the Google Fonts API key programmatically
        
        Args:
            api_key (str): Google Fonts API key
        """
        os.environ['GOOGLE_FONTS_API_KEY'] = api_key

class DocxApp:
    def __init__(self, root):
        self.root = root
        self.root.title("CODEDocx Converter")
        
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        self.root.geometry(f"{screen_width}x{screen_height}")
        
        self.doc = Document()
        self.q_count = 1
        
        # Font selection
        self.selected_font = "Arial"  # Default font
        
        self.create_widgets()
    
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=BOTH, expand=YES)

        style = ttk.Style()
        style.configure('Transparent.TButton', borderwidth=0, background=self.root.cget('background'))
        style.map('Transparent.TButton', background=[('active', self.root.cget('background'))])

        # Create a frame for the top bar
        top_bar = ttk.Frame(main_frame)
        top_bar.pack(fill=X, pady=(0, 10))

        # Add CODEDocx label to the left of the top bar
        logo_label = ttk.Label(top_bar, text="CODEDocx", font=("Arial", 24, "bold"))
        logo_label.pack(side=LEFT, padx=(0, 10))

        # Load and resize the images
        light_icon = Image.open("light_mode.png")
        dark_icon = Image.open("dark_mode.png")
        icon_size = (24, 24)
        light_icon = light_icon.resize(icon_size, Image.LANCZOS)
        dark_icon = dark_icon.resize(icon_size, Image.LANCZOS)
        self.light_icon = ImageTk.PhotoImage(light_icon)
        self.dark_icon = ImageTk.PhotoImage(dark_icon)

        # Add theme toggle button to the right of the top bar
        self.theme_var = ttk.BooleanVar(value=True)
        self.theme_toggle = ttk.Button(
            top_bar,
            image=self.light_icon,
            command=self.toggle_theme,
            style='Transparent.TButton',
            cursor="hand2"
        )
        self.theme_toggle.pack(side=RIGHT)

        # Create content frame
        content_frame = ttk.Frame(main_frame)
        content_frame.pack(fill=BOTH, expand=YES, pady=(0, 10))

        sections = [
            ("Question:", "question_text"),
            ("Code:", "code_text"),
            ("Output:", "output_text")
        ]
        
        for i, (label_text, attr_name) in enumerate(sections):
            frame = ttk.Frame(content_frame)
            frame.pack(fill=BOTH, expand=YES, pady=(0, 10))
            
            ttk.Label(frame, text=label_text).pack(anchor=W, pady=(0, 5))
            text_area = ScrolledText(frame, wrap=WORD, height=6)
            text_area.pack(fill=BOTH, expand=YES)
            setattr(self, attr_name, text_area)

        options_frame = ttk.Frame(main_frame)
        options_frame.pack(fill=X, pady=(0, 10))

        ttk.Label(options_frame, text="Font Size:").pack(side=LEFT, padx=(0, 5))
        self.font_size = ttk.Combobox(options_frame, values=[10, 12, 14, 16, 18, 20], width=5)
        self.font_size.set(12)
        self.font_size.pack(side=LEFT, padx=(0, 20))

        # Add Font selection dropdown
        ttk.Label(options_frame, text="Font:").pack(side=LEFT, padx=(0, 5))
        font_list = GoogleFontManager.get_font_list() or ['Arial', 'Times New Roman', 'Courier']
        self.font_dropdown = ttk.Combobox(
            options_frame, 
            values=font_list, 
            width=20
        )
        self.font_dropdown.set(self.selected_font)
        self.font_dropdown.pack(side=LEFT, padx=(0, 20))
        
        # Bind selection event
        self.font_dropdown.bind('<<ComboboxSelected>>', self.on_font_select)

        self.bold_var = ttk.BooleanVar()
        self.italic_var = ttk.BooleanVar()
        ttk.Checkbutton(options_frame, text="Bold", variable=self.bold_var).pack(side=LEFT, padx=(0, 10))
        ttk.Checkbutton(options_frame, text="Italic", variable=self.italic_var).pack(side=LEFT)

        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=X, pady=(10, 0))

        ttk.Button(
            button_frame,
            text="Add to DOCX",
            command=self.add_to_docx,
            style='primary.TButton',
        ).pack(side=LEFT, padx=(0, 10))

        ttk.Button(
            button_frame,
            text="Start Batch Import",
            command=self.start_batch_import,
            style='primary.TButton',
        ).pack(side=LEFT, padx=(0, 10))

        ttk.Button(
            button_frame,
            text="Save DOCX",
            command=self.save_docx,
            style='primary.TButton',
        ).pack(side=LEFT)

        # Add process indicator
        self.progress_var = ttk.IntVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=X, pady=(10, 0))

        self.status_label = ttk.Label(main_frame, text="")
        self.status_label.pack(pady=(5, 0))

        self.toggle_theme()

    def on_font_select(self, event):
        """Handle font selection"""
        self.selected_font = self.font_dropdown.get()

    def toggle_theme(self):
        if self.theme_var.get():
            self.root.style.theme_use("cyborg")
            self.theme_toggle.configure(image=self.dark_icon)
        else:
            self.root.style.theme_use("morph")
            self.theme_toggle.configure(image=self.light_icon)
    
        # Update button background to match new theme
        new_bg = self.root.cget('background')
        self.root.style.configure('Transparent.TButton', background=new_bg)
        self.root.style.map('Transparent.TButton', background=[('active', new_bg)])
    
        self.theme_var.set(not self.theme_var.get())

    def add_to_docx(self, is_batch=False):
        question = self.question_text.get("1.0", tk.END).strip()
        code = self.code_text.get("1.0", tk.END).strip()
        output = self.output_text.get("1.0", tk.END).strip()
        font_size = int(self.font_size.get())
        bold = self.bold_var.get()
        italic = self.italic_var.get()
        
        if not question or not code:
            if not is_batch:
                messagebox.showerror("Error", "Question and Code fields cannot be empty!")
            return
        
        q_para = self.doc.add_paragraph()
        q_run = q_para.add_run(f"Q{self.q_count}. ")
        q_run.bold = True
        q_run.font.size = Pt(font_size)
        q_run.font.name = self.selected_font
        q_text_run = q_para.add_run(question)
        q_text_run.bold = True
        q_text_run.font.size = Pt(font_size)
        q_text_run.font.name = self.selected_font
        
        code_label_para = self.doc.add_paragraph()
        code_label_run = code_label_para.add_run("Code--")
        code_label_run.bold = True
        code_label_run.font.size = Pt(font_size)
        code_label_run.font.name = self.selected_font
        
        code_para = self.doc.add_paragraph(code)
        for run in code_para.runs:
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic
            run.font.name = self.selected_font
        
        self.doc.add_paragraph()
        
        # Always add "Output--" label
        output_label_para = self.doc.add_paragraph()
        output_label_run = output_label_para.add_run("Output--")
        output_label_run.bold = True
        output_label_run.font.size = Pt(font_size)
        output_label_run.font.name = self.selected_font
        
        if output:
            output_para = self.doc.add_paragraph(output)
            for run in output_para.runs:
                run.font.size = Pt(font_size)
                run.bold = bold
                run.italic = italic
                run.font.name = self.selected_font
        else:
            # Add an empty paragraph if no output
            self.doc.add_paragraph()
        
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        self.q_count += 1
        
        if not is_batch:
            messagebox.showinfo("Success", "Content added to DOCX!")
        
        # Clear input fields
        for widget in [self.question_text, self.code_text, self.output_text]:
            widget.delete("1.0", tk.END)
    
    def save_docx(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file_path:
            self.doc.save(file_path)
            messagebox.showinfo("Success", f"DOCX file saved at {file_path}!")

    def start_batch_import(self):
        directory = filedialog.askdirectory()
        if directory:
            self.batch_importer = CodeFileImporter(self, directory)
            if self.batch_importer.total_files > 0:
                self.batch_import_files()
            else:
                messagebox.showinfo("No Files Found", "No supported files (.c, .cpp, .py, .html) found in the selected directory.")

    def batch_import_files(self):
        if hasattr(self, 'batch_importer'):
            if self.batch_importer.import_next_file():
                self.root.after(100, self.batch_import_files)  # Schedule next file import
            else:
                # No more files to import
                del self.batch_importer
                self.progress_var.set(100)
                self.status_label.config(text="Batch import completed")
                messagebox.showinfo("Batch Import Complete", f"All files have been processed and added to the DOCX.")

class CodeFileImporter:
    def __init__(self, docx_app, directory):
        self.docx_app = docx_app
        self.directory = directory
        self.current_file_number = 1
        self.total_files = len([f for f in os.listdir(directory) if f.endswith(('.c', '.cpp', '.py', '.html'))])

    def import_next_file(self):
        for ext in ['.c', '.cpp', '.py', '.html']:
            filename = f"{self.current_file_number}{ext}"
            filepath = os.path.join(self.directory, filename)
            if os.path.exists(filepath):
                break
        else:
            return False

        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()

            # Extract initial comment
            if filepath.endswith('.html'):
                comment_match = re.match(r'^<!--(.+?)-->', content, re.DOTALL)
                if comment_match:
                    comment = comment_match.group(1).strip()
                    content = re.sub(r'^<!--(.+?)-->', '', content, flags=re.DOTALL).strip()
            else:
                comment_match = re.match(r'^((#|//|/\*).+\n)*', content, re.MULTILINE)
                if comment_match:
                    comment = comment_match.group(0)
                    comment = re.sub(r'^(#|//|\s*\*)\s?', '', comment, flags=re.MULTILINE).strip()
                    content = re.sub(r'^((#|//|/\*).+\n)*', '', content, flags=re.MULTILINE).strip()

            # Set question text
            self.docx_app.question_text.delete('1.0', tk.END)
            self.docx_app.question_text.insert('1.0', comment if comment else "")

            # Set code text
            self.docx_app.code_text.delete('1.0', tk.END)
            self.docx_app.code_text.insert('1.0', content)

            # Clear output text
            self.docx_app.output_text.delete('1.0', tk.END)

            # Add to DOCX
            self.docx_app.add_to_docx(is_batch=True)

            # Update progress
            progress = (self.current_file_number / self.total_files) * 100
            self.docx_app.progress_var.set(progress)
            self.docx_app.status_label.config(text=f"Processing file {self.current_file_number} of {self.total_files}")

            # Increment file number
            self.current_file_number += 1
            return True

        except Exception as e:
            messagebox.showerror("Error", f"Could not process file {filename}: {str(e)}")
            return False

if __name__ == "__main__":
    root = ttk.Window(themename="cyborg")
    app = DocxApp(root)
    root.mainloop()